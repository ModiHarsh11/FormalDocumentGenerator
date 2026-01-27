import os
import re
from datetime import datetime

from django.shortcuts import render
from django.http import HttpResponse
from django.utils import timezone
from django.conf import settings

from .models import DocumentLog
from .constants import DESIGNATION_MAP

# Gemini
import google.generativeai as genai
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# ReportLab
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# DOCX
from docx import Document


# ================== FONT SETUP ==================

FONT_DIR = os.path.join(settings.BASE_DIR, "static", "fonts")

HINDI_REG = os.path.join(FONT_DIR, "NotoSansDevanagari-Regular.ttf")
HINDI_BOLD = os.path.join(FONT_DIR, "NotoSansDevanagari-Bold.ttf")

pdfmetrics.registerFont(TTFont("Hindi", HINDI_REG))
pdfmetrics.registerFont(TTFont("HindiBold", HINDI_BOLD))


# ================== HELPERS ==================

def clean_text(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return text.strip()


# ================== VIEWS ==================

POSITIONS = list(DESIGNATION_MAP.keys())


def home(request):
    return render(request, "generator/index.html", {"positions": POSITIONS})


def generate_document(request):
    if request.method == "POST":
        language = request.POST.get("language")
        ref_id = request.POST.get("reference_id")
        order_date = request.POST.get("order_date") or timezone.now().strftime("%d/%m/%Y")
        prompt = request.POST.get("content")

        from_pos = request.POST.get("from_position")
        to_pos = request.POST.get("to_position")

        if language == "Hindi":
            system_prompt = """
आप एक भारतीय सरकारी कार्यालय आदेश लेखक हैं।
केवल आदेश की मुख्य सामग्री लिखिए।
कोई 
दस्तावेज़ प्रकार, हेडर, संदर्भ संख्या, दिनांक, हस्ताक्षर न लिखें।
पूर्णतः शुद्ध हिंदी में लिखें।
"""
        else:
            system_prompt = """
You are an Indian government office order drafting officer.
Write ONLY the body content.
Do not include document type, header, reference, date, or signature.
"""

        model = genai.GenerativeModel("gemini-2.5-flash-lite")
        response = model.generate_content(system_prompt + prompt)
        content = clean_text(response.text)

        DocumentLog.objects.create(
            document_type="Office Order",
            language=language,
            reference_id=ref_id,
            content=content,
        )

        request.session.update({
            "language": language,
            "reference_id": ref_id,
            "order_date": order_date,
            "content": content,
            "from_pos": from_pos,
            "to_pos": to_pos,
        })

        return render(request, "generator/index.html", {
            "generated_text": content,
            "positions": POSITIONS
        })

    return home(request)


def download_pdf(request):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="BISAG_Office_Order.pdf"'

    doc = SimpleDocTemplate(response, pagesize=A4, leftMargin=40, rightMargin=40, topMargin=40)
    story = []

    lang = request.session["language"]
    content = request.session["content"]
    ref_id = request.session["reference_id"]
    date = request.session["order_date"]

    from_pos = DESIGNATION_MAP[request.session["from_pos"]]["hi" if lang == "Hindi" else "en"]
    to_pos = DESIGNATION_MAP[request.session["to_pos"]]["hi" if lang == "Hindi" else "en"]

    # ---------- HEADER ----------
    header_font = "HindiBold" if lang == "Hindi" else "Helvetica-Bold"

    header = ParagraphStyle("header", alignment=TA_CENTER, fontName=header_font, fontSize=12)

    if lang == "Hindi":
        header_text = """
भास्करचार्य अंतरिक्ष अनुप्रयोग एवं भू-सूचना विज्ञान संस्थान (BISAG-N)<br/>
इलेक्ट्रॉनिक्स एवं सूचना प्रौद्योगिकी मंत्रालय (MeitY)<br/>
भारत सरकार
"""
    else:
        header_text = """
Bhaskaracharya Institute for Space Applications & Geo-Informatics (BISAG-N)<br/>
Ministry of Electronics and Information Technology (MeitY)<br/>
Government of India
"""

    story.append(Paragraph(header_text, header))
    story.append(Spacer(1, 15))

    # ---------- REF & DATE ----------
    
    ref_label = "सं:" if lang == "Hindi" else "Ref. No:"
    date_label = "दिनांक:" if lang == "Hindi" else "Date:"

    meta = Table([
    [f"<b>{ref_label} {ref_id}</b>", f"<b>{date_label} {date}</b>"]
    ], colWidths=[410,110])


    story.append(meta)
    story.append(Spacer(1,20))

    # ---------- DOCUMENT TYPE ----------
    title = "कार्यालय आदेश" if lang == "Hindi" else "OFFICE ORDER"

    title_style = ParagraphStyle(
    "title",
    alignment=TA_CENTER,
    underline=True,
    fontName="HindiBold" if lang == "Hindi" else "Helvetica-Bold",
    fontSize=13,
    )


    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 20))

    # ---------- BODY ----------
    body_style = ParagraphStyle(
        "body",
        fontName="Hindi" if lang == "Hindi" else "Helvetica",
        fontSize=11,
        leading=16
    )

    for para in content.split("\n\n"):
        story.append(Paragraph(para.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 10))

    story.append(Spacer(1, 30))

    # ---------- FROM (RIGHT) ----------
    from_style = ParagraphStyle(
        "from",
        alignment=TA_RIGHT,
        fontName=header_font,
        fontSize=12
    )

    story.append(Paragraph(from_pos, from_style))
    story.append(Spacer(1, 20))

    # ---------- TO (LEFT) ----------
    to_style = ParagraphStyle(
        "to",
        alignment=TA_LEFT,
        fontName=header_font,
        fontSize=12
    )

    to_label = "प्रति:<br/>" if lang == "Hindi" else "To:<br/>"
    story.append(Paragraph(f"{to_label}<br/>{to_pos}, BISAG-N", to_style))

    doc.build(story)
    return response


def download_docx(request):
    lang = request.session["language"]
    content = request.session["content"]
    ref_id = request.session["reference_id"]
    date = request.session["order_date"]

    from_pos = DESIGNATION_MAP[request.session["from_pos"]]["hi" if lang == "Hindi" else "en"]
    to_pos = DESIGNATION_MAP[request.session["to_pos"]]["hi" if lang == "Hindi" else "en"]

    doc = Document()

    # ---------- HEADER ----------
    header_lines = [
        "भास्करचार्य अंतरिक्ष अनुप्रयोग एवं भू-सूचना विज्ञान संस्थान (BISAG-N)",
        "इलेक्ट्रॉनिक्स एवं सूचना प्रौद्योगिकी मंत्रालय (MeitY)",
        "भारत सरकार"
    ] if lang == "Hindi" else [
        "Bhaskaracharya Institute for Space Applications & Geo-Informatics (BISAG-N)",
        "Ministry of Electronics and Information Technology (MeitY)",
        "Government of India"
    ]

    for line in header_lines:
        p = doc.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run(line)
        run.bold = True

    doc.add_paragraph("")

    # ---------- REF & DATE ----------
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = 4500000
    table.columns[1].width = 2000000

    left = table.rows[0].cells[0].paragraphs[0]
    right = table.rows[0].cells[1].paragraphs[0]

    left.add_run(
        ("सं: " if lang == "Hindi" else "Ref. No: ") + ref_id
    ).bold = True

    right.alignment = 2  # RIGHT
    right.add_run(
        ("दिनांक: " if lang == "Hindi" else "Date: ") + date
    ).bold = True

    doc.add_paragraph("")

    # ---------- DOCUMENT TYPE ----------
    title = "कार्यालय आदेश" if lang == "Hindi" else "OFFICE ORDER"
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(title)
    run.bold = True
    run.underline = True

    doc.add_paragraph("")

    # ---------- BODY ----------
    for para in content.split("\n\n"):
        doc.add_paragraph(para)

    doc.add_paragraph("")

    # ---------- FROM (RIGHT) ----------
    p = doc.add_paragraph()
    p.alignment = 2  # RIGHT
    p.add_run(from_pos).bold = True

    doc.add_paragraph("")

    # ---------- TO (LEFT) ----------
    p = doc.add_paragraph()
    p.add_run("प्रति:\n" if lang == "Hindi" else "To:\n").bold = True
    p.add_run(to_pos + ", BISAG-N")

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="BISAG_Office_Order.docx"'
    doc.save(response)
    return response

